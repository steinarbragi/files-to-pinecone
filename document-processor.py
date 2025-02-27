import os
import re
from typing import List, Dict, Any, Tuple
import uuid

# Document processing
import PyPDF2
import docx
import pandas as pd
import openpyxl
from bs4 import BeautifulSoup
import fitz  # PyMuPDF for better PDF handling
import pytesseract
from PIL import Image
import io
import tempfile

# Embedding and vector DB
import pinecone
from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Environment variables (store these in a .env file)
import dotenv
dotenv.load_dotenv()

PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
PINECONE_ENVIRONMENT = os.getenv("PINECONE_ENVIRONMENT")  # e.g., "us-west1-gcp"
PINECONE_INDEX_NAME = os.getenv("PINECONE_INDEX_NAME")

# Configure Tesseract path if not in standard location
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Uncomment and modify for Windows
# For Linux/Mac, ensure tesseract is installed and in path

# Initialize embedding model
embedding_model = SentenceTransformer('all-MiniLM-L6-v2')  # Small, fast model for testing
# For production, consider these stronger models:
# - 'all-mpnet-base-v2' (better quality)
# - 'multi-qa-mpnet-base-dot-v1' (optimized for retrieval)

# Create text splitter for chunking documents
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=512,
    chunk_overlap=50,
    separators=["\n\n", "\n", ". ", " ", ""]
)

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF files using PyMuPDF (fitz) for better handling of
    text, tables, and images with OCR fallback.
    """
    text = ""
    try:
        # Open the PDF with PyMuPDF
        doc = fitz.open(file_path)
        
        for page_num, page in enumerate(doc):
            # Extract text directly - this handles most text including tables better than PyPDF2
            page_text = page.get_text("text")
            text += page_text + "\n"
            
            # Handle images with OCR if text content is minimal
            if len(page_text.strip()) < 100:  # Likely image-heavy page
                images = page.get_images(full=True)
                
                # Process each image on the page
                for img_index, img_info in enumerate(images):
                    xref = img_info[0]  # get the XREF of the image
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # Use OCR to extract text from the image
                    try:
                        image = Image.open(io.BytesIO(image_bytes))
                        img_text = pytesseract.image_to_string(image)
                        if img_text.strip():
                            text += f"[Image Text {page_num+1}-{img_index+1}]: {img_text}\n"
                    except Exception as img_err:
                        print(f"Error processing image in PDF {file_path}, page {page_num+1}: {img_err}")
            
            # Basic table detection - look for grid-like structures
            tables = page.find_tables()
            if tables and tables.tables:
                for i, table in enumerate(tables):
                    text += f"[Table {page_num+1}-{i+1}]:\n"
                    
                    # Convert table to text representation
                    rows = table.extract()
                    for row in rows:
                        text += " | ".join([str(cell) for cell in row]) + "\n"
                    text += "\n"
        
    except Exception as e:
        print(f"Error extracting content from PDF {file_path}: {e}")
        # Fall back to PyPDF2 if PyMuPDF fails
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as fallback_err:
            print(f"Fallback extraction also failed for {file_path}: {fallback_err}")
    
    return text

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX files with improved table handling and image extraction.
    """
    text = ""
    try:
        doc = docx.Document(file_path)
        
        # Process paragraphs
        for para in doc.paragraphs:
            # Skip empty paragraphs
            if para.text.strip():
                text += para.text + "\n"
        
        # Process tables with better formatting
        for table_index, table in enumerate(doc.tables):
            text += f"\n[Table {table_index+1}]:\n"
            
            # Get all cells, find the maximum width for each column
            col_widths = {}
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    content = cell.text.strip()
                    col_widths[i] = max(col_widths.get(i, 0), len(content))
            
            # Format the table with padding for better readability
            for row in table.rows:
                row_text = ""
                for i, cell in enumerate(row.cells):
                    content = cell.text.strip()
                    # Pad content for alignment
                    row_text += content.ljust(col_widths[i] + 2)
                text += row_text + "\n"
            text += "\n"
        
        # Extract images (need to save them temporarily to process with OCR)
        image_index = 0
        temp_dir = tempfile.mkdtemp()
        
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_index += 1
                try:
                    # Save image to temp file
                    image_data = rel.target_part.blob
                    img_path = f"{temp_dir}/image_{image_index}.png"
                    with open(img_path, "wb") as img_file:
                        img_file.write(image_data)
                    
                    # OCR the image
                    img = Image.open(img_path)
                    img_text = pytesseract.image_to_string(img)
                    if img_text.strip():
                        text += f"\n[Image {image_index} Text]: {img_text}\n"
                except Exception as img_err:
                    print(f"Error processing image in DOCX {file_path}: {img_err}")
        
    except Exception as e:
        print(f"Error extracting text from DOCX {file_path}: {e}")
    
    return text

def extract_text_from_xlsx(file_path: str) -> str:
    """
    Extract text from Excel files with improved table structure preservation
    and handling of charts/images where possible.
    """
    text = ""
    try:
        # Using pandas for basic table extraction
        xls = pd.ExcelFile(file_path)
        
        for sheet_name in xls.sheet_names:
            text += f"\n[Sheet: {sheet_name}]\n"
            
            # Read with pandas for structured data
            try:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                if not df.empty:
                    # Format small tables in full
                    if len(df) <= 100:
                        text += df.to_string(index=False) + "\n\n"
                    else:
                        # For large tables, show header, first few rows, and summary
                        text += "Table Headers: " + ", ".join(df.columns.astype(str)) + "\n"
                        text += "First 5 rows:\n" + df.head(5).to_string(index=False) + "\n"
                        text += f"[...{len(df)-10} more rows...]\n"
                        text += "Last 5 rows:\n" + df.tail(5).to_string(index=False) + "\n"
                        
                        # Add statistical summary for numerical columns
                        num_cols = df.select_dtypes(include=['number']).columns
                        if not num_cols.empty:
                            text += "\nNumerical Column Statistics:\n"
                            stats = df[num_cols].describe().to_string()
                            text += stats + "\n"
            except Exception as sheet_err:
                print(f"Error with pandas reading sheet {sheet_name} in {file_path}: {sheet_err}")
            
            # Use openpyxl for additional information (charts, images, formulas)
            try:
                workbook = openpyxl.load_workbook(file_path, data_only=True)
                sheet = workbook[sheet_name]
                
                # Get chart information if available
                if hasattr(sheet, '_charts') and sheet._charts:
                    text += f"\n[Charts in {sheet_name}]: {len(sheet._charts)} charts found\n"
                    for i, chart in enumerate(sheet._charts):
                        chart_title = getattr(chart, 'title', f'Chart {i+1}')
                        text += f"Chart {i+1}: {chart_title}\n"
                
                # Extract comments
                comments = []
                for cell_address, cell in sheet._comment_cache.items() if hasattr(sheet, '_comment_cache') else []:
                    author = getattr(cell, 'author', 'Unknown')
                    content = getattr(cell, 'content', 'No content')
                    comments.append(f"Comment at {cell_address} by {author}: {content}")
                
                if comments:
                    text += "\n[Comments]:\n" + "\n".join(comments) + "\n"
                
                # Check for images/drawings
                if hasattr(sheet, '_images') and sheet._images:
                    text += f"\n[Images in {sheet_name}]: {len(sheet._images)} images found\n"
                
            except Exception as openpyxl_err:
                print(f"Error with openpyxl analysis for {sheet_name} in {file_path}: {openpyxl_err}")
            
            text += "\n" + "-"*40 + "\n"
        
    except Exception as e:
        print(f"Error extracting text from Excel {file_path}: {e}")
    
    return text

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from plain text files."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with different encodings if utf-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            print(f"Error extracting text from TXT {file_path}: {e}")
            return ""
    except Exception as e:
        print(f"Error extracting text from TXT {file_path}: {e}")
        return ""

def extract_text_from_html(file_path: str) -> str:
    """Extract text from HTML files."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file.read(), 'html.parser')
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()
            # Get text
            text = soup.get_text(separator='\n')
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            return text
    except Exception as e:
        print(f"Error extracting text from HTML {file_path}: {e}")
        return ""

def process_document(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Process a document and extract its text based on file extension.
    Returns the extracted text and metadata.
    """
    file_name = os.path.basename(file_path)
    file_ext = os.path.splitext(file_name)[1].lower()
    
    # Extract text based on file extension
    if file_ext == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif file_ext == '.docx':
        text = extract_text_from_docx(file_path)
    elif file_ext in ['.xlsx', '.xls']:
        text = extract_text_from_xlsx(file_path)
    elif file_ext == '.txt':
        text = extract_text_from_txt(file_path)
    elif file_ext in ['.html', '.htm']:
        text = extract_text_from_html(file_path)
    else:
        print(f"Unsupported file type: {file_ext}")
        text = ""
    
    # Create metadata
    metadata = {
        "source": file_path,
        "filename": file_name,
        "filetype": file_ext,
        "created_at": str(os.path.getctime(file_path)),
        "modified_at": str(os.path.getmtime(file_path)),
    }
    
    return text, metadata

def chunk_text(text: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Split text into chunks and attach metadata to each chunk.
    Returns a list of documents with text and metadata.
    """
    chunks = text_splitter.split_text(text)
    documents = []
    
    for i, chunk in enumerate(chunks):
        # Skip empty chunks
        if not chunk.strip():
            continue
            
        # Create a unique ID for each chunk
        chunk_id = str(uuid.uuid4())
        
        # Copy metadata and add chunk-specific info
        chunk_metadata = metadata.copy()
        chunk_metadata.update({
            "chunk_id": chunk_id,
            "chunk_index": i,
            "chunk_count": len(chunks)
        })
        
        documents.append({
            "id": chunk_id,
            "text": chunk,
            "metadata": chunk_metadata
        })
    
    return documents

def get_embeddings(documents: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Generate embeddings for each document chunk.
    Returns documents with embeddings added.
    """
    # Extract text from documents
    texts = [doc["text"] for doc in documents]
    
    # Generate embeddings
    embeddings = embedding_model.encode(texts)
    
    # Add embeddings to documents
    for i, doc in enumerate(documents):
        doc["embedding"] = embeddings[i].tolist()
    
    return documents

def initialize_pinecone():
    """Initialize Pinecone client and ensure index exists."""
    # Initialize Pinecone
    pinecone.init(api_key=PINECONE_API_KEY, environment=PINECONE_ENVIRONMENT)
    
    # Check if index exists
    if PINECONE_INDEX_NAME not in pinecone.list_indexes():
        # Create index if it doesn't exist
        # Adjust dimensions to match your embedding model
        pinecone.create_index(
            name=PINECONE_INDEX_NAME,
            dimension=embedding_model.get_sentence_embedding_dimension(),
            metric="cosine"
        )
    
    # Connect to the index
    index = pinecone.Index(PINECONE_INDEX_NAME)
    return index

def upload_to_pinecone(documents: List[Dict[str, Any]], batch_size: int = 100):
    """
    Upload document embeddings to Pinecone in batches.
    """
    # Initialize Pinecone
    index = initialize_pinecone()
    
    # Prepare vectors in Pinecone format
    vectors = []
    for doc in documents:
        vectors.append({
            "id": doc["id"],
            "values": doc["embedding"],
            "metadata": {
                "text": doc["text"],
                **doc["metadata"]
            }
        })
    
    # Upload vectors in batches
    for i in range(0, len(vectors), batch_size):
        batch = vectors[i:i+batch_size]
        index.upsert(vectors=batch)
        print(f"Uploaded batch {i//batch_size + 1}/{(len(vectors)-1)//batch_size + 1} to Pinecone")
    
    print(f"Successfully uploaded {len(vectors)} vectors to Pinecone index '{PINECONE_INDEX_NAME}'")

def extract_content_from_image(file_path: str) -> str:
    """Extract text content from image files using OCR."""
    text = ""
    try:
        # Open image with PIL
        image = Image.open(file_path)
        
        # Convert image to text using OCR
        text = pytesseract.image_to_string(image)
        
        # Add metadata about the image
        text += f"\n[Image Properties]: Format: {image.format}, Size: {image.size}, Mode: {image.mode}\n"
        
    except Exception as e:
        print(f"Error extracting text from image {file_path}: {e}")
    
    return text

def process_folder(folder_path: str):
    """
    Process all supported documents in a folder and upload to Pinecone.
    """
    # Supported file extensions
    supported_extensions = ['.pdf', '.docx', '.xlsx', '.xls', '.txt', '.html', '.htm', 
                           '.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif']
    
    all_documents = []
    
    # Walk through all files in folder and subfolders
    for root, _, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.join(root, file)
            file_ext = os.path.splitext(file)[1].lower()
            
            if file_ext in supported_extensions:
                print(f"Processing: {file_path}")
                
                # Process images directly
                if file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif']:
                    text = extract_content_from_image(file_path)
                    metadata = {
                        "source": file_path,
                        "filename": os.path.basename(file_path),
                        "filetype": file_ext,
                        "created_at": str(os.path.getctime(file_path)),
                        "modified_at": str(os.path.getmtime(file_path)),
                        "content_type": "image"
                    }
                else:
                    # Extract text and metadata
                    text, metadata = process_document(file_path)
                
                if text:
                    # Chunk the text
                    chunked_docs = chunk_text(text, metadata)
                    
                    # Add to collection
                    all_documents.extend(chunked_docs)
                    
                    print(f"Extracted {len(chunked_docs)} chunks from {file_path}")
                else:
                    print(f"No text extracted from {file_path}")
    
    print(f"Total documents processed: {len(all_documents)}")
    
    # Generate embeddings
    print("Generating embeddings...")
    documents_with_embeddings = get_embeddings(all_documents)
    
    # Upload to Pinecone
    print("Uploading to Pinecone...")
    upload_to_pinecone(documents_with_embeddings)

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Process documents and upload to Pinecone")
    parser.add_argument("folder_path", help="Path to folder containing documents")
    args = parser.parse_args()
    
    process_folder(args.folder_path)
